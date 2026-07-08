from pathlib import Path

from docx import Document
from docx.table import _Cell
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("社交新零售电商平台_2.5用户端功能测试.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LABEL_FILL = "F2F4F7"
WHITE = "FFFFFF"
GRAY = "666666"

# compact_reference_guide preset; table text uses the named dense_test_matrix override.
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
FUNCTION_WIDTHS = [3456, 3888, 2016]
UI_WIDTHS = [7632, 1728]


MODULES = [
    {
        "title": "用户注册与登录测试",
        "description": "验证普通用户注册、账号密码登录、登录态保持及异常登录处理。",
        "entry": "用户端“注册/登录”页面",
        "cases": [
            ("注册时不填写手机号、密码等必填项，直接点击“注册”", "页面逐项提示必填信息，注册请求不提交。"),
            ("输入格式错误的手机号或不符合规则的密码并提交", "页面提示对应格式错误，不能创建账号。"),
            ("使用已注册手机号再次注册", "系统提示账号已存在，不生成重复用户。"),
            ("输入错误账号或密码并点击“登录”", "系统提示账号或密码错误，停留在登录页且不生成登录态。"),
            ("输入正确账号和密码并点击“登录”", "登录成功，保存有效登录态并进入用户端首页。"),
            ("清除或伪造登录凭证后访问个人中心、购物车等受保护页面", "系统拦截访问并引导重新登录。"),
        ],
        "ui": [
            "页面标题、手机号、密码、确认密码等文字及输入框排列正常",
            "密码输入默认脱敏，显示/隐藏密码操作正常",
            "注册、登录及页面切换按钮支持触摸操作且点击区域合理",
            "输入框可正常唤起键盘，手机号输入使用合适的键盘类型",
            "校验错误、登录失败和网络异常提示清晰且不遮挡关键操作",
        ],
    },
    {
        "title": "首页、分类与商品搜索测试",
        "description": "验证首页商品展示、分类/品牌浏览、关键词搜索、筛选排序与分页加载。",
        "entry": "用户端首页、分类页、品牌页及搜索结果页",
        "cases": [
            ("进入首页并下拉刷新", "轮播、分类入口、推荐商品等内容正常加载，刷新后数据同步更新。"),
            ("选择一级/二级分类或品牌", "展示对应范围内的在售商品，不混入不相关或已下架商品。"),
            ("输入存在的商品关键词进行搜索", "返回名称或关键属性匹配的商品，并显示正确的价格、图片和商家信息。"),
            ("输入不存在的关键词进行搜索", "返回空结果状态并提供返回或重新搜索入口，不显示历史脏数据。"),
            ("切换价格、销量等排序条件并连续翻页/上拉加载", "排序规则生效，分页无重复、遗漏或顺序跳变。"),
        ],
        "ui": [
            "首页轮播图、分类图标、商品卡片和底部导航显示完整",
            "搜索框、筛选项和排序状态清晰，选中态可识别",
            "商品图片比例一致，长标题省略合理，价格信息无截断",
            "下拉刷新、上拉加载及页面返回等触摸交互正常",
            "加载中、空数据、图片加载失败和网络异常状态显示正常",
        ],
    },
    {
        "title": "商品详情与SKU选择测试",
        "description": "验证商品详情、SKU规格、库存、评价预览、商家入口及推广码承接。",
        "entry": "商品列表/搜索结果点击商品后进入“商品详情”页面",
        "cases": [
            ("打开正常在售商品详情", "展示商品图片、名称、价格、库存、详情、商家信息、平均评分及最新评价。"),
            ("切换不同SKU规格", "价格、库存、图片等SKU信息同步更新，当前选中规格明确。"),
            ("未选择完整规格时点击“加入购物车”或“立即购买”", "提示先选择完整规格，不进入购物车或结算页。"),
            ("选择库存不足或库存为0的SKU并提交购买", "系统禁止购买并提示库存不足。"),
            ("从带promotionCode的推广链接进入详情页", "推广码被正确读取和缓存；有效推广码与当前商品匹配，无效码不影响普通浏览。"),
            ("点击商家信息或“查看全部评价”", "分别进入正确商家主页或商品评价列表，商品上下文保持正确。"),
        ],
        "ui": [
            "商品轮播图、价格、促销标签、详情图文与评价预览显示正常",
            "SKU弹层层级正确，已选/不可选规格状态区分明显",
            "“加入购物车”“立即购买”等底部按钮固定区域无重叠",
            "长图、长文本和多规格内容可顺畅滚动且不横向溢出",
            "匿名评价隐藏真实昵称和头像，评分星级显示准确",
        ],
    },
    {
        "title": "购物车测试",
        "description": "验证商品加入购物车、数量修改、选择结算、删除及失效商品处理。",
        "entry": "商品详情页“加入购物车”或底部导航“购物车”",
        "cases": [
            ("选择有效SKU和数量后加入购物车", "加入成功，同一SKU按规则合并，购物车角标和数量更新。"),
            ("增加或减少购物车商品数量", "数量在1至可售库存范围内变化，小计和合计金额实时重算。"),
            ("将数量修改为超过库存或非法值", "系统拒绝修改并提示库存限制，原有效数量保持不变。"),
            ("勾选部分商品并点击“去结算”", "仅携带已选且有效的商品进入结算预览。"),
            ("删除单个商品或批量删除已选商品", "二次确认后删除成功，金额、数量和空状态同步更新。"),
            ("购物车中商品下架、SKU失效或库存不足", "商品进入失效区域或被禁止勾选，并给出明确原因。"),
        ],
        "ui": [
            "商品缩略图、规格、单价、数量步进器和选择框排列正常",
            "全选、单选与结算金额的联动状态显示正确",
            "数量加减、删除、滑动操作和底部结算按钮触摸响应正常",
            "失效商品样式与正常商品有明显区分，处理入口清晰",
            "空购物车、加载失败和价格变化提示显示正常",
        ],
    },
    {
        "title": "收货地址测试",
        "description": "验证收货地址新增、编辑、删除、默认地址和结算页地址选择。",
        "entry": "个人中心“收货地址”或订单确认页“选择地址”",
        "cases": [
            ("不填写收货人、手机号或详细地址直接保存", "页面提示必填项，地址不保存。"),
            ("输入错误手机号或超长字段后保存", "系统给出格式/长度提示，不提交非法数据。"),
            ("填写完整有效信息新增地址", "地址保存成功并出现在列表中。"),
            ("编辑已有地址并保存", "修改内容立即生效，其他地址不受影响。"),
            ("将某地址设为默认地址", "该地址成为唯一默认地址，结算页优先选中。"),
            ("删除当前默认地址", "二次确认后删除成功，系统按规则选择新的默认地址或显示未设置状态。"),
        ],
        "ui": [
            "地址列表中的收货人、手机号、地区、详细地址和默认标识显示完整",
            "新增/编辑表单字段顺序合理，地区选择器可正常展开和回填",
            "默认地址开关、编辑、删除和选择操作触摸响应正常",
            "删除操作有明确二次确认，成功/失败提示位置合理",
            "无地址时的空状态与“新增地址”入口清晰可见",
        ],
    },
    {
        "title": "订单确认与优惠计算测试",
        "description": "验证立即购买/购物车结算预览、地址与配送方式、优惠券、积分及金额计算。",
        "entry": "购物车“去结算”或商品详情“立即购买”进入订单确认页",
        "cases": [
            ("未选择收货地址且选择配送方式时提交订单", "系统提示选择收货地址，不创建订单。"),
            ("切换配送/自提方式并选择自提点", "配送费、自提信息和可填写字段按方式正确变化。"),
            ("选择可用优惠券、满减活动或积分抵扣", "优惠条件校验通过，优惠明细和应付金额按规则实时重算。"),
            ("选择不满足门槛、已过期或不适用商品的优惠券", "优惠券不可选或提示不可用原因，不产生错误抵扣。"),
            ("提交订单前商品价格、库存或活动发生变化", "系统重新校验并提示变化，使用最新有效数据确认后方可下单。"),
            ("携带有效推广码提交订单", "订单记录正确的推广归因；无效或不匹配推广码不产生归因。"),
            ("连续快速点击“提交订单”", "系统防止重复提交，仅生成一个有效订单并正确扣减库存。"),
        ],
        "ui": [
            "地址、商品、配送方式、优惠明细、积分和订单金额层级清晰",
            "优惠券/自提点选择弹层内容完整，选中态明确",
            "商品金额、运费、优惠额和应付金额的币种与小数位显示一致",
            "提交按钮在加载期间禁用并显示进度，避免重复操作",
            "金额变化、库存不足和优惠失效提示醒目且可理解",
        ],
    },
    {
        "title": "订单支付测试",
        "description": "验证支付请求创建、支付宝沙箱/模拟支付、支付结果查询及异常恢复。",
        "entry": "订单提交成功页或待支付订单详情页“立即支付”",
        "cases": [
            ("对有效待支付订单发起支付", "创建与订单金额一致的支付请求并进入支付流程。"),
            ("支付成功并返回用户端", "支付状态更新为成功，订单进入待发货/待处理状态且不可重复支付。"),
            ("支付失败或用户取消支付", "订单保持待支付状态，记录失败原因并允许重新发起支付。"),
            ("支付过程中网络中断，随后重新进入订单详情", "系统主动查询支付结果并展示最终一致状态。"),
            ("对已支付、已取消或已关闭订单再次发起支付", "系统拒绝支付并提示当前订单状态不允许支付。"),
            ("支付回调重复到达", "系统幂等处理，不重复记账、不重复扣库存或增加积分。"),
        ],
        "ui": [
            "订单号、应付金额、支付方式和支付倒计时显示清晰",
            "支付确认按钮醒目，点击后有加载态并防止重复触发",
            "支付成功、失败、取消和处理中状态使用明确图标与文字",
            "从第三方支付页面返回后的页面跳转和返回按钮行为正常",
            "支付异常时提供重试、查看订单等可恢复操作",
        ],
    },
    {
        "title": "订单查询与状态流转测试",
        "description": "验证订单列表/详情、状态筛选、取消订单、确认收货及订单状态同步。",
        "entry": "个人中心“我的订单”",
        "cases": [
            ("进入全部订单并按待付款、待发货、待收货、已完成筛选", "列表仅展示对应状态订单，分页加载完整且顺序正确。"),
            ("打开任一订单详情", "订单商品、金额、地址/自提、支付、配送和状态信息与订单一致。"),
            ("取消允许取消的待支付订单", "二次确认后订单变为已取消，库存/优惠资源按规则释放。"),
            ("尝试取消已发货、已完成等不可取消订单", "系统拒绝操作并提示当前状态不可取消。"),
            ("对待收货订单点击“确认收货”", "确认后订单变为已完成，并开放评价/售后等符合规则的入口。"),
            ("商家更新发货状态后刷新订单", "用户端及时获取最新状态，列表和详情状态一致。"),
        ],
        "ui": [
            "订单卡片中的状态、商品、数量、金额和操作按钮排列正常",
            "状态筛选标签可横向滚动或完整显示，选中态明确",
            "不同订单状态仅显示允许的操作按钮，不出现越权入口",
            "取消、确认收货等关键操作均有二次确认",
            "列表空状态、加载更多和状态刷新提示显示正常",
        ],
    },
    {
        "title": "售后服务测试",
        "description": "验证仅退款/退货退款申请、凭证上传、进度查询、补充凭证与取消申请。",
        "entry": "订单详情/订单商品“申请售后”或个人中心“售后记录”",
        "cases": [
            ("对不满足售后条件的订单或商品发起申请", "系统拒绝申请并提示不满足售后条件。"),
            ("选择仅退款，填写有效原因和金额后提交", "申请创建成功，金额不超过可退金额，状态为待审核。"),
            ("选择退货退款并填写退货相关信息", "申请创建成功，后续按审核、退货、收货、退款流程流转。"),
            ("上传不支持格式、超大小或超数量的凭证图片", "系统拒绝不合规图片并提示限制；合法图片可预览和删除。"),
            ("在允许状态下补充凭证或取消售后申请", "信息追加/取消成功，列表和详情状态同步。"),
            ("商家审核通过/拒绝或退款完成后查看详情", "处理意见、退款金额、时间线和最终状态准确展示。"),
        ],
        "ui": [
            "售后类型、原因、金额、说明和凭证上传区域排列清晰",
            "图片选择、预览、删除和上传进度显示正常",
            "售后状态时间线顺序正确，当前节点突出显示",
            "可执行操作随状态变化正确显示，禁用态清晰",
            "审核拒绝原因、退款失败等异常信息完整且可阅读",
        ],
    },
    {
        "title": "商品评价测试",
        "description": "验证已完成订单评价、评分与图文内容、匿名展示、评价查询及删除。",
        "entry": "已完成订单“去评价”、个人中心“我的评价”或商品详情评价区",
        "cases": [
            ("对未完成订单、非本人订单或已评价订单再次评价", "系统拒绝提交并提示当前订单商品不可评价。"),
            ("评分为空/越界或评价内容不符合规则时提交", "页面提示校验错误，评价不保存。"),
            ("填写1至5星、文字和合法图片后提交评价", "评价创建成功并关联正确订单明细、商品和SKU。"),
            ("勾选匿名后提交评价", "评价参与评分统计，但商品页隐藏真实昵称、头像及用户标识。"),
            ("按评分或有图条件查看商品全部评价", "筛选和分页结果正确，平均分与有效评价数量一致。"),
            ("删除本人的评价", "评价从我的评价和商品展示中移除，平均评分/评价数按有效数据重算。"),
        ],
        "ui": [
            "评分星级、评价输入框、匿名选项及图片上传区显示正常",
            "星级选择反馈即时，已选与未选状态区分明显",
            "评价图片网格、预览和删除操作触摸响应正常",
            "长评价内容换行合理，匿名标识和时间显示清晰",
            "筛选项、分页加载和无评价状态显示正常",
        ],
    },
    {
        "title": "优惠券、满减与秒杀测试",
        "description": "验证优惠券领取/查看、满减规则展示、秒杀活动及资格校验。",
        "entry": "首页活动入口、商家主页、商品详情、领券中心及“我的优惠券”",
        "cases": [
            ("领取未开始、已结束、已领完或超过限领次数的优惠券", "系统禁止领取并提示具体原因。"),
            ("领取有效优惠券", "领取成功并出现在“我的优惠券”，库存和领取次数正确扣减。"),
            ("查看可用、已使用、已过期优惠券", "分类及状态正确，适用范围、门槛和有效期完整展示。"),
            ("进入满减活动商品并达到/未达到门槛", "详情和结算页正确展示规则，只有满足条件时产生优惠。"),
            ("进入当前秒杀活动并校验商品资格", "展示正确秒杀价、活动时间、剩余库存和购买限制。"),
            ("秒杀开始前、结束后或超限购买", "系统拒绝下单并提示活动状态或限购原因。"),
        ],
        "ui": [
            "优惠券面额、门槛、有效期、适用范围和状态标签显示完整",
            "已领取、已抢光、未开始等按钮状态差异明确",
            "秒杀倒计时及时刷新，结束后页面状态自动切换",
            "活动价、原价和优惠标签层级清晰且不产生价格歧义",
            "活动空状态、库存不足和领取失败提示显示正常",
        ],
    },
    {
        "title": "会员与积分测试",
        "description": "验证会员等级、成长值、积分余额/明细、签到、积分抵扣和兑换优惠券。",
        "entry": "个人中心“会员中心/积分中心”",
        "cases": [
            ("进入会员中心", "展示当前等级、成长值、下一等级条件及对应会员权益。"),
            ("查询积分余额、积分明细和成长值明细", "余额与明细汇总一致，来源、变动值和时间准确。"),
            ("当日首次签到", "签到成功并按规则增加积分；当日再次签到不重复奖励。"),
            ("在结算页输入可用积分试算抵扣", "抵扣上限、换算规则和应付金额计算正确。"),
            ("积分不足或兑换已失效/库存不足的优惠券", "系统拒绝兑换，不扣积分并提示原因。"),
            ("使用足够积分兑换有效优惠券", "积分扣减与优惠券发放保持一致，明细各生成一条记录。"),
        ],
        "ui": [
            "会员等级、成长进度、积分余额和权益卡片显示正常",
            "成长进度条比例与数值一致，不超出容器",
            "签到、兑换、明细筛选等按钮触摸响应正常",
            "积分增减使用清晰的正负号、颜色和来源描述",
            "兑换确认、积分不足和重复签到提示显示正常",
        ],
    },
    {
        "title": "分销推广测试",
        "description": "验证分销员申请、推广商品、推广链接/二维码、归因、佣金与提现。",
        "entry": "个人中心“分销中心/推广中心”",
        "cases": [
            ("普通用户未提交完整实名、身份证或银行卡资料即申请分销员", "页面提示补全必填资料，申请不提交。"),
            ("提交完整有效资料申请分销员", "申请成功并显示待审核；审核结果在状态页和消息中同步。"),
            ("未通过审核的用户访问分销工作台或生成推广链接", "系统拒绝访问或操作，并提示当前审核状态。"),
            ("已通过用户选择开放分销的商品生成推广链接/二维码", "生成的地址指向用户端商品详情并携带正确promotionCode。"),
            ("被推广用户通过链接下单并完成订单", "系统按推广码和商品完成归因，生成正确佣金明细。"),
            ("推广订单取消、退款或售后完成", "佣金按规则取消、回退或保持冻结，统计金额同步更新。"),
            ("可提现佣金不足、银行卡无效或金额超限时申请提现", "系统拒绝申请且不扣可提现余额；合法申请生成提现记录。"),
        ],
        "ui": [
            "申请表单、证件图片和银行卡信息输入区域层级清晰",
            "审核中、已通过、已拒绝等状态及原因显示完整",
            "推广商品、二维码、复制链接和启停推广操作正常",
            "累计、冻结、可提现佣金数值和佣金明细展示一致",
            "敏感身份及银行卡信息按规则脱敏显示",
        ],
    },
    {
        "title": "个人中心与消息通知测试",
        "description": "验证个人资料、手机号/密码修改、功能入口及订单/售后/活动消息的查询和已读处理。",
        "entry": "底部导航“我的/个人中心”及“消息中心”",
        "cases": [
            ("进入个人中心", "展示当前用户头像、昵称、手机号脱敏信息及各业务入口。"),
            ("修改昵称、头像等允许编辑的个人资料", "保存成功后页面即时更新，重新登录后仍保持。"),
            ("修改手机号时输入非法号码或已绑定号码", "系统拒绝修改并提示原因；合法号码修改成功。"),
            ("原密码错误、两次新密码不一致或新密码不合规时修改密码", "系统拒绝修改并给出对应提示。"),
            ("查询订单、审核、售后、活动和系统消息", "消息可按类型正确筛选，内容、关联业务和时间准确。"),
            ("打开未读消息或执行全部已读", "未读状态和角标及时更新，刷新后保持一致。"),
            ("点击带业务关联的消息", "跳转到对应订单、售后、活动或审核结果详情，不跳错对象。"),
        ],
        "ui": [
            "头像、昵称、会员/积分信息和功能宫格排列正常",
            "手机号等敏感信息脱敏显示，密码输入框默认隐藏内容",
            "消息类型标签、未读角标、时间和摘要显示清晰",
            "消息列表滑动、筛选、单条已读和全部已读操作正常",
            "无消息、加载失败和跳转目标失效时有明确反馈",
        ],
    },
]


def set_run_font(run, ascii_name="Calibri", east_asia="微软雅黑", size=None, bold=None, color=None):
    run.font.name = ascii_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7C3D0", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        grid_index = 0
        for tc in row._tr.tc_lst:
            tc_pr = tc.get_or_add_tcPr()
            grid_span = tc_pr.first_child_found_in("w:gridSpan")
            span = int(grid_span.get(qn("w:val"))) if grid_span is not None else 1
            width = sum(widths[grid_index:grid_index + span])
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell = _Cell(tc, table)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            grid_index += span


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, *, bold=False, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    return p


def add_function_table(doc, module):
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    set_table_borders(table)

    set_cell_text(table.cell(0, 0), "功能描述", bold=True, color=DARK_BLUE)
    desc = table.cell(0, 1).merge(table.cell(0, 2))
    set_cell_text(desc, module["description"])
    set_cell_shading(table.cell(0, 0), LABEL_FILL)

    set_cell_text(table.cell(1, 0), "功能入口", bold=True, color=DARK_BLUE)
    entry = table.cell(1, 1).merge(table.cell(1, 2))
    set_cell_text(entry, module["entry"])
    set_cell_shading(table.cell(1, 0), LABEL_FILL)

    headers = ["输入/动作", "期望的输出/响应", "实际的输出/响应"]
    for i, text in enumerate(headers):
        set_cell_text(table.cell(2, i), text, bold=True, color=DARK_BLUE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(2, i), HEADER_FILL)
    repeat_table_header(table.rows[2])

    for action, expected in module["cases"]:
        cells = table.add_row().cells
        set_cell_text(cells[0], action)
        set_cell_text(cells[1], expected)
        set_cell_text(cells[2], "待测试", color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
        prevent_row_split(table.rows[-1])

    set_table_geometry(table, FUNCTION_WIDTHS)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_ui_table(doc, items):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    set_cell_text(table.cell(0, 0), "检查项", bold=True, color=DARK_BLUE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), "是否正常", bold=True, color=DARK_BLUE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(0, 0), HEADER_FILL)
    set_cell_shading(table.cell(0, 1), HEADER_FILL)
    repeat_table_header(table.rows[0])
    for item in items:
        cells = table.add_row().cells
        set_cell_text(cells[0], item)
        set_cell_text(cells[1], "待测试", color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, UI_WIDTHS)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=GRAY)


def build():
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("社交新零售电商平台｜用户端测试方案")
    set_run_font(hr, size=9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    kr = kicker.add_run("测试方案 · 用户端")
    set_run_font(kr, size=11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    tr = title.add_run("2.5 功能/用户界面测试要点")
    set_run_font(tr, size=24, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("社交新零售电商平台——用户端功能测试")
    set_run_font(sr, size=14, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    mr1 = meta.add_run("测试对象：")
    set_run_font(mr1, size=10, bold=True, color=GRAY)
    mr2 = meta.add_run("uni-app 用户端　　")
    set_run_font(mr2, size=10, color=GRAY)
    mr3 = meta.add_run("测试类型：")
    set_run_font(mr3, size=10, bold=True, color=GRAY)
    mr4 = meta.add_run("黑盒功能测试 / 用户界面测试")
    set_run_font(mr4, size=10, color=GRAY)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(10)
    nr1 = note.add_run("编写说明：")
    set_run_font(nr1, size=10, bold=True, color=DARK_BLUE)
    nr2 = note.add_run(
        "本文档参照《智慧考勤系统》测试方案的章节和双表结构编写。"
        "“实际的输出/响应”和“是否正常”栏在执行测试前统一标记为“待测试”，测试完成后填写实际结果。"
    )
    set_run_font(nr2, size=10, color="333333")

    add_heading(doc, "2.5 功能/用户界面测试要点", 1)

    for index, module in enumerate(MODULES, start=1):
        add_heading(doc, f"2.5.{index} {module['title']}", 2)
        add_heading(doc, f"2.5.{index}.1 功能测试", 3)
        add_function_table(doc, module)
        add_heading(doc, f"2.5.{index}.2 用户界面测试", 3)
        add_ui_table(doc, module["ui"])

    doc.core_properties.title = "社交新零售电商平台 2.5 用户端功能测试"
    doc.core_properties.subject = "功能/用户界面测试要点"
    doc.core_properties.author = "项目测试组"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
